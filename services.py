
from fastapi import HTTPException
from docx import Document
import logging
from google import genai
from google.genai import types
from io import BytesIO

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    client = genai.Client()
except Exception as e:
    # If the API key isn't set, this will fail. We'll handle it below.
    client = None
    print(f"Warning: Gemini Client failed to initialize. Please set GEMINI_API_KEY. Error: {e}")

# --- Core AI Function ---

async def translate_and_format_pdf_with_gemini(content: bytes, filename: str) -> str:
    """
    Uses Gemini-2.5-Flash to perform OCR, translation, and formatting in one call.
    
    Args:
        content: The byte content of the PDF file.
        filename: The original file name.

    Returns:
        The formatted English text as a string.
    """
    if client is None:
        raise HTTPException(
            status_code=503, 
            detail="Gemini API Client is not initialized. Please ensure GEMINI_API_KEY environment variable is set."
        )

    # System instruction defines the model's persona and rules for the task
    system_instruction = (
        "You are an expert legal document translator and formatter. "
        "Your task is to analyze the provided scanned PDF (written in Bangla), "
        "convert ALL text to professional, clear English, and meticulously preserve the "
        "original document's structure, formatting, and layout. "
        "Use Markdown syntax to represent headings, lists and paragraphs exactly "
        "Ignore the bangla stamp and tables"
        "For legal documents, maintain fidelity to the original sections and line breaks."
    )
    
    # User prompt requests the specific action
    prompt = (
        f"The attached file, named '{filename}', is a scanned legal case file written in Bengali (Bangla). "
        "Translate the entire document content into English. "
        "Maintain the original formatting and section layout as closely as possible using Markdown. "
        "Start with a suitable title for the translated document."
    )

    config = types.GenerateContentConfig(
        system_instruction=system_instruction
    )

    try:
        # The multimodal call (text + image/PDF)
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                types.Part.from_bytes(
                    data=content,
                    mime_type='application/pdf',
                ),
                prompt
            ],
            config=config
        )
        return response.text
        
    except Exception as e:
        # Catch any API-related errors
        raise HTTPException(status_code=500, detail=f"AI processing failed: {e}")
        
# --- DOCX Text Extraction Helper ---

def extract_text_from_docx(content: bytes) -> str:
    """
    Extracts structured text content (paragraphs, headings, and lists) 
    from a DOCX file's byte content and converts it into a simple Markdown template.
    Tables are intentionally skipped per user request.
    """
    text = []
    try:
        # Load DOCX from a byte stream
        document = Document(BytesIO(content))

        # 1. Process Paragraphs (Headings and Lists)
        for paragraph in document.paragraphs:
            stripped_text = paragraph.text.strip()
            if not stripped_text:
                continue

            style_name = paragraph.style.name
            
            # Check for Headings
            if style_name.startswith('Heading'):
                # Use Markdown headers to preserve structure in the text template
                level = int(style_name.split()[-1])
                text.append(f"{'#' * level} {stripped_text}")
            
            # Check for Lists (by style name, often starting with 'List')
            elif style_name.startswith('List'):
                # Default to bullet, use '1. ' if it looks like a numbered list
                prefix = '* '
                if 'Number' in style_name or 'num' in style_name.lower():
                    prefix = '1. ' 
                
                text.append(f"{prefix}{stripped_text}")
                
            # Default paragraph
            else:
                text.append(stripped_text)
        
        # Tables are intentionally skipped here based on user request.
        
        return "\n".join(text)
    except Exception as e:
        print(f"Error extracting structured text from DOCX: {e}")
        # Return an empty string on failure
        return ""
# --- Core AI Function 2 (Refinement) ---

async def refine_english_markdown(markdown_text: str, sample_text_content: str = None) -> str:
    """
    Uses a second Gemini-2.5-Flash call to clean up, standardize, and finalize
    the translated English text and Markdown structure, optionally using a sample text template for style.
    
    Args:
        markdown_text: The translated and initially formatted Markdown text.
        sample_text_content: Optional string containing the extracted text template from a sample DOCX.

    Returns:
        The cleaned, refined Markdown text as a string.
    """
    if client is None:
        return markdown_text # Skip refinement if client is not initialized

    # Build the contents list dynamically
    contents = []
    text_instruction = ""

    # If a sample style text is provided, include it in the prompt
    if sample_text_content:
        # Add instruction that the preceding part is the style sample
        text_instruction = (
            "The following section is a TEXT TEMPLATE from a desired style reference document. "
            "Analyze its structure, headings, and legal phrasing. You MUST use this structure and style "
            "for the final output of the translated case file.\n\n"
            f"--- START OF STYLE TEMPLATE ---\n{sample_text_content}\n--- END OF STYLE TEMPLATE ---\n\n"
        )
        print(f"DEBUG: Using {len(sample_text_content)} characters of DOCX content as a style template.")


    # System instruction for refinement
    system_instruction = (
        "You are a professional editor specializing in legal document standardization. "
        "Your task is to proofread, correct grammatical errors, and ensure consistent "
        "legal terminology in the provided translated text. "
        "Crucially, standardize the Markdown usage (headings, lists, paragraphs) according to the provided "
        "style template (if present), and remove any extraneous introductory/closing phrases or junk text, "
        "outputting ONLY the clean, finalized legal document content in Markdown format."
    )
    
    # User prompt for refinement
    text_prompt = (
        text_instruction + 
        "Refine the following translated legal document text. Ensure all formatting is strictly consistent, "
        "legal terminology is correct, and grammar is flawless. Return ONLY the final, cleaned Markdown content:\n\n"
        f"--- START OF DOCUMENT TO REFINE ---\n{markdown_text}"
    )
    
    # Add the main text prompt
    contents.append(text_prompt)

    # Define the generation configuration, including the system instruction
    config = types.GenerateContentConfig(
        system_instruction=system_instruction
    )

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=contents, # Use the dynamic contents list
            config=config
        )
        # Return the refined text, or the original if the response is empty
        return response.text if response.text else markdown_text
        
    except Exception as e:
        print(f"Warning: Refinement AI call failed: {e}. Using original translated text.")
        return markdown_text # Return original text on failure to keep the process moving
        
# --- DOCX Generation Helper ---

def generate_docx_from_markdown(markdown_text: str) -> BytesIO:
    """
    Converts simple markdown text (headers, paragraphs) into a DOCX file buffer.
    This is a simplification; for complex markdown, you'd use a dedicated parser.
    """
    doc = Document()
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check for headers (simple # detection)
        if line.startswith('### '):
            doc.add_heading(line.lstrip('### ').strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.lstrip('## ').strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.lstrip('# ').strip(), level=1)
        
        # Check for lists (simple * or - detection)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line.lstrip('* ').lstrip('- ').strip(), style='List Bullet')
        
        # Default paragraph
        else:
            doc.add_paragraph(line)

    # Save the document to an in-memory byte buffer
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer